"""
Generate BTP2_REPORT.docx with formatting matching the PDF exactly.
A4, Times New Roman 12pt, 1.5 line spacing, justified body text.
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

MD_FILE  = "BTP2_REPORT.md"
OUT_FILE = "BTP2_REPORT.docx"
FIGS_DIR = "figures"

# ── Colours ──────────────────────────────────────────────────────────────────
C_HDR_BG  = RGBColor(0xD0, 0xD0, 0xD0)
C_ALT_ROW = RGBColor(0xF5, 0xF5, 0xF5)
C_BDR     = RGBColor(0x33, 0x33, 0x33)
C_BLACK   = RGBColor(0x00, 0x00, 0x00)

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# A4 page size
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(3.17)

# Page numbers bottom-right
def add_page_numbers(section):
    footer = section.footer
    para   = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)
    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = para.add_run()
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld2)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

add_page_numbers(section)

# ── Style helpers ─────────────────────────────────────────────────────────────
TNR = 'Times New Roman'

def _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=8,
                  line_spacing=1.5, keep_with_next=False,
                  page_break_before=False, left_indent=0, hanging=0):
    pf = para.paragraph_format
    pf.alignment         = align
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line_spacing
    pf.keep_with_next    = keep_with_next
    pf.page_break_before = page_break_before
    if left_indent:
        pf.left_indent = Pt(left_indent)
    if hanging:
        pf.first_line_indent = Pt(-hanging)
        pf.left_indent       = Pt(hanging)

def _run(para, text, bold=False, italic=False, size=12, color=None, name=TNR):
    run = para.add_run(text)
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def add_normal(text, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               size=12, space_before=0, space_after=8):
    para = doc.add_paragraph()
    _set_para_fmt(para, align=align, space_before=space_before, space_after=space_after)
    _add_inline(para, text, size=size, bold=bold, italic=italic)
    return para

def _add_inline(para, text, size=12, bold=False, italic=False):
    """Parse **bold**, *italic*, `code`, and plain text into runs."""
    # pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _run(para, text[pos:m.start()], bold=bold, italic=italic, size=size)
        raw = m.group(0)
        if raw.startswith('**'):
            _run(para, m.group(2), bold=True, italic=italic, size=size)
        elif raw.startswith('*'):
            _run(para, m.group(3), bold=bold, italic=True, size=size)
        else:  # backtick code
            _run(para, m.group(4), bold=bold, italic=italic, size=size,
                 name='Courier New')
        pos = m.end()
    if pos < len(text):
        _run(para, text[pos:], bold=bold, italic=italic, size=size)

def add_heading_chapter(text, centered=False):
    """h2 = chapter heading: 14pt bold uppercase, page break before."""
    para = doc.add_paragraph()
    _set_para_fmt(para,
                  align=WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=18,
                  line_spacing=1.0, keep_with_next=True,
                  page_break_before=True)
    _run(para, text.upper(), bold=True, size=14)
    return para

def add_heading_section(text, level=3):
    """h3/h4/h5 section headings."""
    size   = 12
    italic = (level == 5)
    para   = doc.add_paragraph()
    _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=12, space_after=6,
                  line_spacing=1.0, keep_with_next=True)
    _run(para, text, bold=True, italic=italic, size=size)
    return para

def add_caption(text):
    """Figure/Table caption: 10pt italic centered."""
    para = doc.add_paragraph()
    _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=2, space_after=12, line_spacing=1.0)
    _run(para, text, italic=True, size=10)
    return para

def add_bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=4, line_spacing=1.5)
    _add_inline(para, text)
    for run in para.runs:
        run.font.name = TNR
        run.font.size = Pt(12)
    return para

def add_numbered(text, n):
    para = doc.add_paragraph(style='List Number')
    _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=4, line_spacing=1.5)
    _add_inline(para, text)
    for run in para.runs:
        run.font.name = TNR
        run.font.size = Pt(12)
    return para

def add_spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        _set_para_fmt(p, space_before=0, space_after=0, line_spacing=1.0)

def add_image(path, width_cm=14.0):
    abs_path = os.path.join(FIGS_DIR, os.path.basename(path))
    if os.path.exists(abs_path):
        para = doc.add_paragraph()
        _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=6, space_after=4, line_spacing=1.0)
        run = para.add_run()
        run.add_picture(abs_path, width=Cm(width_cm))
    else:
        para = add_normal(f'[Figure: {path}]', italic=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    return para

# ── Table helper ─────────────────────────────────────────────────────────────
def _set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '%02X%02X%02X' % (rgb[0], rgb[1], rgb[2]))
    tcPr.append(shd)

def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '333333')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_table_from_rows(header_row, data_rows, col_widths=None):
    """Build a formatted table from header + data rows."""
    ncols = len(header_row)
    nrows = len(data_rows)
    tbl   = doc.add_table(rows=nrows + 1, cols=ncols)
    tbl.style = 'Table Grid'

    # Default equal column widths
    total_w = Cm(14.96)  # page width minus margins
    if col_widths:
        widths = [Cm(w) for w in col_widths]
    else:
        w = total_w / ncols
        widths = [w] * ncols

    for ci in range(ncols):
        for ri in range(nrows + 1):
            tbl.cell(ri, ci).width = widths[ci]

    # Header row
    hdr = tbl.rows[0]
    for ci, htext in enumerate(header_row):
        cell = hdr.cells[ci]
        _set_cell_bg(cell, C_HDR_BG)
        _set_cell_border(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(htext)
        run.font.name   = TNR
        run.font.size   = Pt(10)
        run.font.bold   = True
        run.font.color.rgb = C_BLACK

    # Data rows
    for ri, row in enumerate(data_rows):
        bg = C_ALT_ROW if ri % 2 == 1 else RGBColor(0xFF, 0xFF, 0xFF)
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            _set_cell_bg(cell, bg)
            _set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline(para, val, size=10)
            for run in para.runs:
                run.font.name = TNR
                run.font.size = Pt(10)

    # Space after table
    p = doc.add_paragraph()
    _set_para_fmt(p, space_before=0, space_after=6, line_spacing=1.0)
    return tbl

# ── Markdown parser ───────────────────────────────────────────────────────────
def parse_table_line(line):
    """Extract cell values from a markdown table row."""
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells

def is_table_separator(line):
    stripped = line.strip().strip('|')
    return bool(re.match(r'^[\s\-|:]+$', stripped))

def parse_md(md_text):  # noqa: C901
    lines = md_text.split('\n')
    i = 0
    in_title_page  = False
    in_right_div   = False
    in_blockquote  = False
    table_buffer   = []
    list_buffer    = []
    ol_buffer      = []
    in_neko_answer = False  # track the verbatim KGMiner answer block

    def flush_list():
        for item in list_buffer:
            add_bullet(item)
        list_buffer.clear()

    def flush_ol():
        for n, item in enumerate(ol_buffer, 1):
            add_numbered(item, n)
        ol_buffer.clear()

    def flush_table():
        if not table_buffer:
            return
        header = table_buffer[0]
        data   = [r for r in table_buffer[2:] if r]  # skip separator row
        add_table_from_rows(header, data)
        table_buffer.clear()

    while i < len(lines):
        line = lines[i]
        raw  = line

        # ── Title page HTML ──────────────────────────────────────────────────
        if '<div class="titlepage">' in line:
            in_title_page = True
            i += 1
            continue
        if in_title_page:
            if '</div>' in line:
                in_title_page = False
                # page break already implied by next heading
                i += 1
                continue
            # h1
            m = re.match(r'\s*<h1>(.*?)</h1>', line)
            if m:
                para = doc.add_paragraph()
                _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                              space_before=30, space_after=24, line_spacing=1.4)
                _run(para, m.group(1), bold=True, size=18)
                i += 1; continue
            # img in title page
            m = re.match(r'\s*<p><img src="([^"]+)"', line)
            if m:
                add_spacer()
                add_image(m.group(1), width_cm=4.5)
                i += 1; continue
            # p with em/strong
            m = re.match(r'\s*<p>(.*?)</p>', line, re.DOTALL)
            if m:
                inner = m.group(1)
                # Handle <br> as newline
                parts = re.split(r'<br\s*/?>', inner)
                for pi, part in enumerate(parts):
                    part = re.sub(r'<[^>]+>', '', part).strip()
                    if not part:
                        continue
                    is_bold   = bool(re.search(r'<strong>', inner))
                    is_italic = bool(re.search(r'<em>', inner)) and not is_bold
                    para = doc.add_paragraph()
                    _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_before=3, space_after=3, line_spacing=1.6)
                    clean = re.sub(r'<[^>]+>', '', part)
                    _run(para, clean, bold=is_bold, italic=is_italic, size=12)
            i += 1; continue

        # ── Right-aligned div (signatures) ────────────────────────────────────
        if re.match(r'\s*<div style="text-align:right', line):
            in_right_div = True
            i += 1; continue
        if in_right_div:
            if '</div>' in line:
                in_right_div = False
                i += 1; continue
            m = re.match(r'\s*<p[^>]*>(.*?)</p>', line)
            if m:
                inner = m.group(1)
                clean = re.sub(r'<[^>]+>', '', inner)
                # parse inline bold
                para = doc.add_paragraph()
                _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.RIGHT,
                              space_before=2, space_after=2, line_spacing=1.0)
                is_bold = '<strong>' in inner
                _run(para, clean, bold=is_bold, size=11)
            i += 1; continue

        # ── h2 with front-heading class ────────────────────────────────────────
        m = re.match(r'\s*<h2 class="front-heading">(.*?)</h2>', line)
        if m:
            flush_list(); flush_ol(); flush_table()
            add_heading_chapter(m.group(1), centered=True)
            i += 1; continue

        # ── Skip remaining HTML tags ───────────────────────────────────────────
        if re.match(r'\s*<[^>]+>', line) and not line.strip().startswith('!'):
            # &nbsp; lines
            if '&nbsp;' in line or line.strip() in ('&nbsp;', ''):
                add_spacer()
            i += 1; continue

        # ── Table rows ────────────────────────────────────────────────────────
        if re.match(r'\s*\|', line):
            flush_list(); flush_ol()
            cells = parse_table_line(line)
            table_buffer.append(cells)
            i += 1; continue
        else:
            if table_buffer:
                flush_table()

        # ── Headings ──────────────────────────────────────────────────────────
        m = re.match(r'^(#{2,6})\s+(.*)', line)
        if m:
            flush_list(); flush_ol()
            level = len(m.group(1))
            text  = m.group(2).strip()
            if level == 2:
                add_heading_chapter(text, centered=False)
            elif level == 3:
                add_heading_section(text, level=3)
            elif level == 4:
                add_heading_section(text, level=4)
            else:
                add_heading_section(text, level=5)
            i += 1; continue

        # ── Horizontal rule ────────────────────────────────────────────────────
        if re.match(r'^---+\s*$', line):
            flush_list(); flush_ol()
            para = doc.add_paragraph()
            _set_para_fmt(para, space_before=6, space_after=6, line_spacing=1.0)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '777777')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1; continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith('>'):
            flush_list(); flush_ol()
            text = line.lstrip('> ').strip()
            para = doc.add_paragraph()
            _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=4, line_spacing=1.5,
                          left_indent=24)
            _add_inline(para, text, italic=True)
            i += 1; continue

        # ── Bullet list ────────────────────────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            flush_ol()
            list_buffer.append(m.group(2))
            i += 1; continue
        else:
            if list_buffer:
                flush_list()

        # ── Numbered list ─────────────────────────────────────────────────────
        m = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if m:
            flush_list()
            ol_buffer.append(m.group(2))
            i += 1; continue
        else:
            if ol_buffer:
                flush_ol()

        # ── Image ─────────────────────────────────────────────────────────────
        m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if m:
            flush_list(); flush_ol()
            add_image(m.group(2), width_cm=13.5)
            i += 1; continue

        # ── Italic-only line = figure/table caption ────────────────────────────
        m = re.match(r'^\*([^*].*[^*])\*\s*$', line.strip())
        if m:
            flush_list(); flush_ol()
            add_caption(m.group(1))
            i += 1; continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            flush_list(); flush_ol()
            i += 1; continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        flush_list(); flush_ol()
        # Check for &nbsp;
        if line.strip() == '&nbsp;':
            add_spacer()
            i += 1; continue

        text = line.strip()
        # Skip pure HTML comment or tag lines we haven't handled
        if re.match(r'^<[^>]+>$', text):
            i += 1; continue

        # Detect if this is a references section paragraph (hanging indent)
        is_ref = bool(re.match(r'^\[\d+\]', text))
        para = doc.add_paragraph()
        if is_ref:
            _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=4,
                          line_spacing=1.5, hanging=24)
            _add_inline(para, text, size=11)
        else:
            _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=8, line_spacing=1.5)
            _add_inline(para, text, size=12)

        i += 1

    flush_list(); flush_ol(); flush_table()

# ── Run ───────────────────────────────────────────────────────────────────────
with open(MD_FILE, encoding='utf-8') as f:
    md = f.read()

parse_md(md)

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
